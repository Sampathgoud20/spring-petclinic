# Create a DOCX version of the user's resume with the added Internship section
from docx import Document

doc = Document()

doc.add_heading('SAMPATH GANAGANI', level=0)
doc.add_paragraph('Email: sampathgoud.ganagani@gmail.com')
doc.add_paragraph('Phone: +91 9392894866')
doc.add_paragraph('LinkedIn: linkedin.com/in/sampath-goud-20a045335')

doc.add_heading('Objective', level=1)
doc.add_paragraph(
"Motivated and detail-oriented DevOps Engineer (Fresher) with hands-on knowledge of CI/CD pipelines, "
"cloud platforms (AWS & Azure), Docker, Kubernetes, Linux, and Infrastructure as Code. Seeking an "
"opportunity to contribute to real-world production systems while continuously learning and implementing "
"DevOps best practices for scalability, reliability, and security."
)

doc.add_heading('Education', level=1)
doc.add_paragraph("Avanthi Institute of Engineering and Technology – B.Tech in Computer Science (Nov 2021 – June 2025)")
doc.add_paragraph("Shivani Junior College – Intermediate MPC (July 2019 – March 2021)")
doc.add_paragraph("MRG KKR High School – SSC (June 2018 – April 2019)")

doc.add_heading('Technical Skills', level=1)
doc.add_paragraph("DevOps & CI/CD: Jenkins Pipelines, Git, GitHub")
doc.add_paragraph("Cloud Platforms: AWS (EC2, S3, IAM, VPC), Microsoft Azure (VMs, Storage, Resource Groups – basic)")
doc.add_paragraph("Containers & Orchestration: Docker, Kubernetes (Pods, Deployments, Services, ConfigMaps, Secrets)")
doc.add_paragraph("Infrastructure as Code: Terraform (basic provisioning)")
doc.add_paragraph("Operating Systems: Linux (Ubuntu, Amazon Linux – commands, shell basics)")
doc.add_paragraph("Monitoring & Reliability: CloudWatch, basic system monitoring concepts, SRE fundamentals")
doc.add_paragraph("Other Tools: YAML, Bash scripting (basic)")

doc.add_heading('Internship Experience', level=1)
doc.add_paragraph("DevOps Intern – Magneq Software (Aug 2025 – Jan 2026)")
intern_points = [
"Built CI/CD pipelines using Jenkins to automate application build and deployment.",
"Containerized applications using Docker and deployed them on Kubernetes clusters.",
"Integrated SonarQube for code quality analysis in the CI pipeline.",
"Used Git and GitHub for version control and team collaboration."
]
for p in intern_points:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Project Experience', level=1)
doc.add_paragraph("Containerized Application Deployment using Docker & Kubernetes")
proj1 = [
"Built Docker images and managed containers using Docker.",
"Deployed applications on Kubernetes cluster using Deployments and Services.",
"Worked with Pods, ReplicaSets, ConfigMaps, and Secrets.",
"Practiced rolling updates and scaling strategies."
]
for p in proj1:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph("CI/CD Pipeline Automation using Jenkins")
proj2 = [
"Designed and implemented a CI/CD pipeline using Jenkins for automated build and deployment.",
"Integrated GitHub repository with Jenkins.",
"Automated application deployment on Linux server.",
"Improved deployment consistency and reduced manual effort."
]
for p in proj2:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph("Cloud Infrastructure Setup (AWS & Azure – Practice)")
proj3 = [
"Launched EC2 instances and Azure VMs.",
"Configured IAM users, security groups, and storage.",
"Monitored resources using CloudWatch.",
"Practiced basic cloud cost-awareness and security concepts."
]
for p in proj3:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Certifications & Co-Curricular Activities', level=1)
doc.add_paragraph("TCS iON Career Edge – Young Professional (Certification)")
doc.add_paragraph("Attended Seminar on Cloud Computing & Personality Development")
doc.add_paragraph("Participated in Big Business Bureau Webinar")

doc.add_heading('Extra-Curricular Activities', level=1)
doc.add_paragraph("NSS Member – Participated in tree plantation and community service initiatives.")

doc.add_heading('Personal Details', level=1)
details = [
"Date of Birth: 13 December 2004",
"Gender: Male",
"Father’s Name: G. Venkanna",
"Mother’s Name: G. Sujatha",
"Languages Known: Telugu, English",
"Hobbies: Travelling, Reading, Watching Movies"
]
for d in details:
    doc.add_paragraph(d)

doc.add_heading('Declaration', level=1)
doc.add_paragraph(
"I hereby declare that the above information is true and correct to the best of my knowledge and belief."
)

path = "/mnt/data/sampath_devops_resume.docx"
doc.save(path)

path